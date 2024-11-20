from transformers import pipeline
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import re
import time
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from config import MOODLE_USERNAME, MOODLE_PASSWORD
from selenium.webdriver.chrome.options import Options
import torch
from gpt4all import GPT4All
import requests
import os

class MoodleAutomation:
    def __init__(self):
        """Initialisation du driver et du modèle IA"""
        self.setup_driver()
        try:
            model_path = os.path.join(os.path.dirname(__file__), 'models', 'mistral-7b-instruct-v0.1.Q4_0.gguf')
            
            if not os.path.exists(model_path):
                print(f"Erreur: Modèle non trouvé à {model_path}")
                raise FileNotFoundError(f"Le modèle n'existe pas à {model_path}")
                
            self.model = GPT4All(model_path)
            print("Modèle IA chargé avec succès")
            
        except Exception as e:
            print(f"Erreur lors du chargement du modèle IA: {e}")

    def setup_driver(self):
        """Initialisation du driver en mode navigation privée"""
        try:
            # Configuration des options Chrome
            chrome_options = Options()
            chrome_options.add_argument("--incognito")
            chrome_options.add_argument("--start-maximized")
            chrome_options.add_argument("--disable-extensions")
            chrome_options.add_argument("--disable-notifications")
            
            # Initialisation du driver avec les options
            self.driver = webdriver.Chrome(options=chrome_options)
            print("Driver initialisé en mode navigation privée")
            
        except Exception as e:
            print(f"Erreur lors de l'initialisation du driver: {e}")
            raise e
        self.sections = [
            "INTRODUCTION", "MOTS CLE", "DONNEES", "CONTEXTE",
            "PROBLEMATIQUE", "CONTRAINTE", "LIVRABLE", "GENERALISATION",
            "PISTES DE SOLUTIONS", "PARTIE THEORIQUE", "PARTIE PRATIQUE",
            "CONCLUSION", "AAV", "BIBLIOGRAPHIE"
        ]

    def connexion_moodle(self, username, password):
        """Connexion via ENT et accès au prosit"""
        try:
            print("1. Accès à la page de connexion Moodle...")
            self.driver.get("https://moodle.cesi.fr/login/index.php")
            time.sleep(2)
            
            # Cliquer sur le bouton ENT
            ent_button = self.driver.find_element(By.CSS_SELECTOR, "a[href*='authCAS=CAS']")
            ent_button.click()
            time.sleep(2)
            
            # Remplir l'identifiant
            login_field = self.driver.find_element(By.ID, "login")
            login_field.send_keys(username)
            time.sleep(1)
            
            submit_button = self.driver.find_element(By.ID, "submit")
            submit_button.click()
            time.sleep(3)
            
            # Remplir le mot de passe
            password_field = self.driver.find_element(By.ID, "passwordInput")
            password_field.send_keys(password)
            time.sleep(1)
            
            sign_in_button = self.driver.find_element(By.ID, "submitButton")
            sign_in_button.click()
            time.sleep(5)
            
            # Accéder au prosit
            print("6. Accès au prosit...")
            self.driver.get("https://moodle.cesi.fr/pluginfile.php/152051/mod_resource/content/4/co/_4_-_Prosit_A_points.html")
            time.sleep(3)   
            
            return True
            
        except Exception as e:
            print(f"Erreur de connexion: {e}")
            return False

    def extraire_prosit(self, url):
        """Extrait le contenu du prosit"""
        try:
            print(f"Accès au prosit: {url}")
            self.driver.get(url)
            
            wait = WebDriverWait(self.driver, 10)
            content = wait.until(
                EC.presence_of_element_located((By.CLASS_NAME, "hBk"))
            )
            
            soup = BeautifulSoup(self.driver.page_source, 'html.parser')
            prosit_content = soup.find('div', class_='hBk desc')
            
            return prosit_content.get_text() if prosit_content else None
            
        except Exception as e:
            print(f"Erreur d'extraction: {e}")
            return None

    def generer_document(self, titre: str, contexte: str, reponses: dict):
        """Génère le document Word avec le contenu structuré"""
        try:
            # Créer le dossier 'resultats' s'il n'existe pas
            if not os.path.exists('resultats'):
                os.makedirs('resultats')
            
            # Créer un nom de fichier basé sur la date et le titre
            date = time.strftime("%Y%m%d-%H%M%S")
            nom_fichier = f"resultats/prosit_{date}.docx"
            
            # Créer le document
            doc = Document()
            
            # Ajouter le titre du prosit
            doc.add_heading(f'Prosit: {titre}', 0)
            
            # Mapping des sections du prosit vers les sections du document
            mapping_sections = {
                "INTRODUCTION": ["generalites"],
                "MOTS CLE": ["mots_cles"],
                "CONTEXTE": [contexte],
                "CONTRAINTE": ["contraintes"],
                "PROBLEMATIQUE": ["problematique"],
                "PISTES DE SOLUTIONS": ["pistes_solution", "hypotheses"],
                "PARTIE THEORIQUE": ["plan_action"],
                "PARTIE PRATIQUE": [],
                "CONCLUSION": [],
                "AAV": [],
                "BIBLIOGRAPHIE": []
            }
            
            # Ajouter chaque section dans l'ordre
            for section in self.sections:
                doc.add_heading(section, level=1)
                
                # Ajouter le contenu correspondant
                if section in mapping_sections:
                    for sous_section in mapping_sections[section]:
                        if sous_section in reponses:
                            doc.add_paragraph(reponses[sous_section])
                        elif sous_section == contexte:
                            doc.add_paragraph(contexte)
            
            # Ajouter un saut de page après chaque section principale
            doc.add_page_break()
            
            # Sauvegarder le document
            doc.save(nom_fichier)
            print(f"\nDocument sauvegardé dans: {nom_fichier}")
            
        except Exception as e:
            print(f"Erreur lors de la génération du document: {e}")

    def completer_prosit(self):
        """Complète automatiquement toutes les sections du prosit"""
        try:
            # Extraire le titre et le contexte
            titre = self.driver.find_element(By.TAG_NAME, "h1").text
            contexte = self.extraire_contexte()
            
            # Dictionnaire pour stocker les réponses
            reponses = {}
            
            # Liste des sections à compléter
            sections = [
                "generalites",
                "mots_cles",
                "contraintes",
                "problematique",
                "hypotheses",
                "pistes_solution",
                "plan_action"
            ]

            # Compléter chaque section
            for section in sections:
                reponse = self.generer_reponse(contexte, section)
                reponses[section] = reponse
                self.completer_section(section, reponse)
                print(f"Section {section} complétée")

            # Générer le document Word
            self.generer_document(titre, contexte, reponses)

        except Exception as e:
            print(f"Erreur lors de la complétion du prosit: {e}")

    def extraire_contexte(self) -> str:
        """Extrait le sujet du prosit et génère un contexte adaptatif"""
        try:
            print("Extraction du sujet du prosit...")
            # Extraire le contenu complet de la page
            contenu_page = self.driver.find_element(By.TAG_NAME, "body").text
            titre = self.driver.find_element(By.TAG_NAME, "h1").text
            
            # Prompt générique adaptable
            prompt = f"""Analyse ce sujet de prosit :

            Titre: {titre}
            Contenu: {contenu_page}

            En tant qu'expert, génère un contexte professionnel qui :
            1. Identifie le domaine technique concerné
            2. Explique la situation professionnelle
            3. Décrit les enjeux principaux
            4. Présente les aspects techniques à considérer
            
            Format: Un paragraphe clair et structuré qui pose le cadre professionnel du problème."""
            
            print("Génération du contexte...")
            contexte = ""
            for token in self.model.generate(
                prompt=prompt,
                max_tokens=300,
                temp=0.7,
                top_k=40,
                top_p=0.9,
                streaming=True
            ):
                contexte += token
            
            return contexte.strip()
                
        except Exception as e:
            print(f"Erreur lors de la génération du contexte: {e}")
            return "Erreur de génération du contexte"

    def completer_section(self, type_section: str, contexte: str):
        """Complète une section spécifique du prosit"""
        try:
            # Générer le contenu
            contenu = self.generer_reponse(contexte, type_section)
            
            # Trouver le champ à remplir
            champ = self.trouver_champ_section(type_section)
            
            if champ:
                # Effacer le contenu existant
                champ.clear()
                # Remplir avec le nouveau contenu
                champ.send_keys(contenu)
                time.sleep(1)  # Attendre que le texte soit bien inséré
            
        except Exception as e:
            print(f"Erreur lors de la complétion de la section {type_section}: {e}")

    def trouver_champ_section(self, type_section: str):
        """Trouve le champ de saisie pour une section donnée"""
        mapping_sections = {
            "generalites": "//textarea[contains(@id, 'generalites')]",
            "mots_cles": "//textarea[contains(@id, 'mots-cles')]",
            "contraintes": "//textarea[contains(@id, 'contraintes')]",
            "problematique": "//textarea[contains(@id, 'problematique')]",
            "hypotheses": "//textarea[contains(@id, 'hypotheses')]",
            "pistes_solution": "//textarea[contains(@id, 'pistes')]",
            "plan_action": "//textarea[contains(@id, 'plan-action')]"
        }
        
        try:
            xpath = mapping_sections.get(type_section)
            return self.driver.find_element(By.XPATH, xpath)
        except:
            print(f"Champ non trouvé pour la section {type_section}")
            return None

    def generer_reponse(self, contexte: str, type_section: str) -> str:
        """Génère une réponse avec GPT4All"""
        try:
            prompt = self.creer_prompt(contexte, type_section)
            print(f"\nGénération pour {type_section}...")
            print(f"Contexte utilisé: {contexte[:100]}...")
            
            # Collecter tous les tokens dans une chaîne
            reponse = ""
            for token in self.model.generate(
                prompt=prompt,
                max_tokens=200,
                temp=0.7,
                top_k=40,
                top_p=0.9,
                streaming=True
            ):
                reponse += token
            
            print(f"Réponse générée: {reponse[:100]}...")
            return self.formater_reponse(reponse, type_section)
            
        except Exception as e:
            print(f"Erreur détaillée de génération pour {type_section}: {str(e)}")
            return f"Erreur de génération pour {type_section}"

    def creer_prompt(self, contexte: str, type_section: str) -> str:
        """Crée un prompt générique adapté selon la section"""
        prompts = {
            "generalites": f"""Analyse ce contexte et liste les points généraux importants à comprendre:
            {contexte}
            Format: Liste de 3-4 points clés.""",
            
            "mots_cles": f"""Identifie les termes techniques et concepts importants dans ce contexte:
            {contexte}
            Format: 5-8 mots-clés séparés par des virgules.""",
            
            "contraintes": f"""Liste les principales contraintes et limitations à considérer:
            {contexte}
            Format: Liste de 3-4 contraintes principales.""",
            
            "problematique": f"""Formule une problématique claire et précise basée sur:
            {contexte}
            Format: Une phrase interrogative qui cible le cœur du problème.""",
            
            "hypotheses": f"""Propose des hypothèses de solutions pour:
            {contexte}
            Format: 2-3 hypothèses de travail réalistes.""",
            
            "pistes_solution": f"""Suggère des solutions concrètes pour:
            {contexte}
            Format: Liste de 3-4 solutions techniques détaillées.""",
            
            "plan_action": f"""Établis un plan d'action structuré pour:
            {contexte}
            Format: Liste numérotée de 4-5 étapes clés."""
        }
        return prompts.get(type_section, "")

    def formater_reponse(self, reponse: str, type_section: str) -> str:
        """Formate la réponse selon le type de section"""
        formatters = {
            "mots_cles": lambda r: ", ".join(r.split()[:10]),
            "contraintes": lambda r: "\n".join([f"- {ligne.strip()}" 
                                              for ligne in r.split(".") 
                                              if ligne.strip()]),
            "plan_action": lambda r: "\n".join([f"{i+1}. {ligne.strip()}" 
                                              for i, ligne in enumerate(r.split(".")) 
                                              if ligne.strip()])
        }
        
        return formatters.get(type_section, lambda r: r.strip())(reponse)

def download_model():
    # Créer le dossier models s'il n'existe pas
    if not os.path.exists('models'):
        os.makedirs('models')
    
    # URL du modèle
    url = "https://gpt4all.io/models/ggml-gpt4all-j-v1.3-groovy.bin"
    
    # Chemin de destination
    model_path = os.path.join('models', 'ggml-gpt4all-j-v1.3-groovy.bin')
    
    print("Téléchargement du modèle en cours...")
    response = requests.get(url, stream=True)
    total_size = int(response.headers.get('content-length', 0))
    
    # Télécharger le fichier
    with open(model_path, 'wb') as file:
        for data in response.iter_content(chunk_size=1024):
            file.write(data)
    
    print("Téléchargement terminé!")

def main():
    url_prosit = "https://moodle.cesi.fr/pluginfile.php/152051/mod_resource/content/4/co/_4_-_Prosit_A_points.html"
    
    automation = MoodleAutomation()
    automation.setup_driver()
    
    try:
        if automation.connexion_moodle(MOODLE_USERNAME, MOODLE_PASSWORD):
            automation.completer_prosit()
    except Exception as e:
        print(f"Erreur: {e}")
    finally:
        input("Appuyez sur Entrée pour fermer...")
        automation.driver.quit()

if __name__ == "__main__":
    main()

