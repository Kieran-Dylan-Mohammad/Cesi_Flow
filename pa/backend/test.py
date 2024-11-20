from transformers import pipeline
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import re
import time
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from config import MOODLE_USERNAME, MOODLE_PASSWORD
from selenium.webdriver.chrome.options import Options
import torch

class PrositAnalyzer:
    def __init__(self):
        # Initialiser le modèle de classification de texte
        self.classifier = pipeline("text-classification", model="facebook/bart-large-mnli")
        self.summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
        
    def analyser_texte(self, texte):
        """Analyse le texte pour extraire les différentes parties"""
        resultats = {}
        
        # Définir les catégories à identifier
        categories = {
            "contexte": "Ce texte décrit-il le contexte?",
            "problematique": "Ce texte pose-t-il une question ou un problème?",
            "contraintes": "Ce texte décrit-il des contraintes ou limitations?",
            "livrables": "Ce texte mentionne-t-il des livrables ou résultats attendus?"
        }
        
        # Diviser le texte en paragraphes
        paragraphes = re.split(r'\n\n|\.\s', texte)
        
        for categorie, prompt in categories.items():
            resultats[categorie] = []
            for para in paragraphes:
                if len(para.strip()) < 10:  # Ignorer les paragraphes trop courts
                    continue
                    
                # Classifier le paragraphe
                prediction = self.classifier(
                    para,
                    candidate_labels=[prompt, "non pertinent"],
                    hypothesis_template="Ce texte est à propos de {}."
                )
                
                if prediction[0]['labels'][0] == prompt and prediction[0]['scores'][0] > 0.7:
                    resultats[categorie].append(para)
        
        # Résumer chaque section
        for categorie in resultats:
            if resultats[categorie]:
                texte_complet = " ".join(resultats[categorie])
                if len(texte_complet) > 100:  # Ne résumer que si assez long
                    resume = self.summarizer(texte_complet, max_length=130, min_length=30)
                    resultats[categorie] = resume[0]['summary_text']
                else:
                    resultats[categorie] = texte_complet
                    
        return resultats

class MoodleAutomation:
    def __init__(self):
        """Initialisation du driver et du modèle IA"""
        self.setup_driver()
        try:
            self.generator = pipeline('text-generation', 
                                   model='gpt2-french', 
                                   device=0 if torch.cuda.is_available() else -1)
            print("Modèle IA chargé avec succès")
        except Exception as e:
            print(f"Erreur lors du chargement du modèle IA: {e}")

    def setup_driver(self):
        """Initialisation du driver en mode navigation privée"""
        try:
            # Configuration des options Chrome
            chrome_options = Options()
            chrome_options.add_argument("--incognito")  # Mode navigation privée
            chrome_options.add_argument("--start-maximized")  # Fenêtre maximisée
            chrome_options.add_argument("--disable-extensions")  # Désactive les extensions
            chrome_options.add_argument("--disable-notifications")  # Désactive les notifications
            
            # Initialisation du driver avec les options
            self.driver = webdriver.Chrome(options=chrome_options)
            print("Driver initialisé en mode navigation privée")
            
        except Exception as e:
            print(f"Erreur lors de l'initialisation du driver: {e}")
            raise e
        self.analyzer = PrositAnalyzer()
        self.sections = [
            "INTRODUCTION", "MOTS CLE", "DONNEES", "CONTEXTE",
            "PROBLEMATIQUE", "CONTRAINTE", "LIVRABLE", "GENERALISATION",
            "PISTES DE SOLUTIONS", "PARTIE THEORIQUE", "PARTIE PRATIQUE",
            "CONCLUSION", "AAV", "BIBLIOGRAPHIE"
        ]

    def connexion_moodle(self):
        """Connexion via ENT et accès au prosit"""
        try:
            print("1. Accès à la page de connexion Moodle...")
            self.driver.get("https://moodle.cesi.fr/login/index.php")
            time.sleep(2)
            
            # 1. Cliquer sur le bouton ENT
            ent_button = self.driver.find_element(By.CSS_SELECTOR, "a[href*='authCAS=CAS']")
            ent_button.click()
            time.sleep(2)
            
            # 2. Remplir l'identifiant et valider
            login_field = self.driver.find_element(By.ID, "login")
            login_field.send_keys(MOODLE_USERNAME)
            time.sleep(1)
            
            submit_button = self.driver.find_element(By.ID, "submit")
            submit_button.click()
            time.sleep(3)
            
            # 3. Remplir le mot de passe et valider
            password_field = self.driver.find_element(By.ID, "passwordInput")
            password_field.send_keys(MOODLE_PASSWORD)
            time.sleep(1)
            
            sign_in_button = self.driver.find_element(By.ID, "submitButton")
            sign_in_button.click()
            time.sleep(5)
            
            # Accéder au prosit
            print("6. Accès au prosit...")
            self.driver.get("https://moodle.cesi.fr/pluginfile.php/152051/mod_resource/content/4/co/_4_-_Prosit_A_points.html")
            time.sleep(3)
            
                
        except Exception as e:
            print(f"Erreur lors de la connexion: {e}")
            return False

    def extraire_prosit(self, url):
        """Extrait le contenu du prosit"""
        try:
            print(f"Accès au prosit: {url}")
            self.driver.get(url)
            
            wait = WebDriverWait(self.driver, 10)
            content = wait.until(
                EC.presence_of_element_located((By.CLASS_NAME, "hBk"))
            )
            
            soup = BeautifulSoup(self.driver.page_source, 'html.parser')
            prosit_content = soup.find('div', class_='hBk desc')
            
            return prosit_content.get_text() if prosit_content else None
            
        except Exception as e:
            print(f"Erreur d'extraction: {e}")
            return None

    def generer_document(self, contenu, nom_fichier="prosit_analyse.docx"):
        """Génère le document Word avec le contenu extrait"""
        doc = Document()
        
        for section in self.sections:
            doc.add_heading(section, level=1)
            doc.add_paragraph("T")  # Placeholder
            doc.add_paragraph()
            
        doc.save(nom_fichier)
        print(f"Document généré: {nom_fichier}")

    def completer_prosit(self):
        """Complète automatiquement toutes les sections du prosit"""
        try:
            # Attendre que la page soit chargée
            WebDriverWait(self.driver, 10).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )

            # Extraire le contexte général
            contexte = self.extraire_contexte()

            # Liste des sections à compléter
            sections = [
                "generalites",
                "mots_cles",
                "contraintes",
                "problematique",
                "hypotheses",
                "pistes_solution",
                "plan_action"
            ]

            # Compléter chaque section
            for section in sections:
                self.completer_section(section, contexte)
                print(f"Section {section} complétée")

        except Exception as e:
            print(f"Erreur lors de la complétion du prosit: {e}")

    def extraire_contexte(self) -> str:
        """Extrait le contexte général du prosit"""
        try:
            # Chercher les éléments contenant le contexte
            elements_contexte = self.driver.find_elements(
                By.XPATH, 
                "//*[contains(text(), 'Contexte') or contains(text(), 'contexte')]"
            )
            return " ".join([elem.text for elem in elements_contexte])
        except Exception as e:
            print(f"Erreur lors de l'extraction du contexte: {e}")
            return ""

    def completer_section(self, type_section: str, contexte: str):
        """Complète une section spécifique du prosit"""
        try:
            # Générer le contenu
            contenu = self.generer_reponse(contexte, type_section)
            
            # Trouver le champ à remplir
            champ = self.trouver_champ_section(type_section)
            
            if champ:
                # Effacer le contenu existant
                champ.clear()
                # Remplir avec le nouveau contenu
                champ.send_keys(contenu)
                time.sleep(1)  # Attendre que le texte soit bien inséré
            
        except Exception as e:
            print(f"Erreur lors de la complétion de la section {type_section}: {e}")

    def trouver_champ_section(self, type_section: str):
        """Trouve le champ de saisie pour une section donnée"""
        mapping_sections = {
            "generalites": "//textarea[contains(@id, 'generalites')]",
            "mots_cles": "//textarea[contains(@id, 'mots-cles')]",
            "contraintes": "//textarea[contains(@id, 'contraintes')]",
            "problematique": "//textarea[contains(@id, 'problematique')]",
            "hypotheses": "//textarea[contains(@id, 'hypotheses')]",
            "pistes_solution": "//textarea[contains(@id, 'pistes')]",
            "plan_action": "//textarea[contains(@id, 'plan-action')]"
        }
        
        try:
            xpath = mapping_sections.get(type_section)
            return self.driver.find_element(By.XPATH, xpath)
        except:
            print(f"Champ non trouvé pour la section {type_section}")
            return None

    def generer_reponse(self, contexte: str, type_section: str) -> str:
        """Génère une réponse adaptée selon le contexte et le type de section"""
        prompts = {
            "generalites": f"Dans le contexte technique suivant: {contexte}, décris les aspects généraux importants.",
            "mots_cles": f"Liste les mots clés techniques pour: {contexte}",
            "contraintes": f"Identifie les contraintes techniques dans: {contexte}",
            "problematique": f"Formule une problématique technique pour: {contexte}",
            "hypotheses": f"Propose des hypothèses techniques pour: {contexte}",
            "pistes_solution": f"Suggère des solutions techniques pour: {contexte}",
            "plan_action": f"Établis un plan d'action technique pour: {contexte}"
        }

        try:
            prompt = prompts.get(type_section, "")
            reponse = self.generator(prompt, 
                                   max_length=200, 
                                   num_return_sequences=1)[0]['generated_text']
            return self.formater_reponse(reponse, type_section)
        except Exception as e:
            print(f"Erreur de génération pour {type_section}: {e}")
            return "Erreur de génération"

    def formater_reponse(self, reponse: str, type_section: str) -> str:
        """Formate la réponse selon le type de section"""
        formatters = {
            "mots_cles": lambda r: ", ".join(r.split()[:10]),
            "contraintes": lambda r: "\n".join([f"- {ligne.strip()}" 
                                              for ligne in r.split(".") 
                                              if ligne.strip()]),
            "plan_action": lambda r: "\n".join([f"{i+1}. {ligne.strip()}" 
                                              for i, ligne in enumerate(r.split(".")) 
                                              if ligne.strip()])
        }
        
        return formatters.get(type_section, lambda r: r.strip())(reponse)

def main():
    url_prosit = "https://moodle.cesi.fr/pluginfile.php/152051/mod_resource/content/4/co/_4_-_Prosit_A_points.html"
    
    automation = MoodleAutomation()
    automation.setup_driver()
    
    try:
        if automation.connexion_moodle():
            automation.completer_prosit()
    except Exception as e:
        print(f"Erreur: {e}")
    finally:
        input("Appuyez sur Entrée pour fermer...")
        automation.driver.quit()

if __name__ == "__main__":
    main()

